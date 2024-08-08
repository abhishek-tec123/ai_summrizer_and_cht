# from flask import Blueprint, request, jsonify
# import io
# import os
# import logging
# import requests
# import PyPDF2
# import unicodedata
# from docx import Document
# import google.generativeai as genai
# import time
# from google.api_core.exceptions import ResourceExhausted


# # Blueprint for document routes
# document_bp = Blueprint('document_bp', __name__)

# # Load environment variables
# api_key = os.getenv("GOOGLE_API_KEY")
# genai.configure(api_key=api_key)

# def extract_text_from_pdf(pdf_stream: io.BytesIO) -> str:
#     logging.info("Extracting text from the PDF.")
#     pdf_reader = PyPDF2.PdfReader(pdf_stream)
#     text = ""
#     for page in pdf_reader.pages:
#         page_text = page.extract_text() or ""
#         text += page_text
#     return text

# def normalize_text(text: str) -> str:
#     return unicodedata.normalize('NFKC', text)

# def chunk_text(text: str, chunk_size: int = 5000) -> list:
#     text = normalize_text(text)
#     chunks = []
#     start = 0
#     while start < len(text):
#         end = start + chunk_size
#         if end < len(text):
#             while end < len(text) and text[end] in ('\uD800', '\uDBFF'):
#                 end += 1
#         chunks.append(text[start:end])
#         start = end
#     return chunks

# def summarize_chunk(llm, chunk: str, retries: int = 3, delay: int = 1) -> str:
#     prompt = f"Please provide a well-defined summary of the following text:\n TEXT: {chunk}"
#     for attempt in range(retries):
#         try:
#             response = llm.generate_content(prompt)
#             summary_text = response.text
#             return summary_text
#         except ResourceExhausted as e:
#             logging.warning(f"Resource exhausted error: {e}. Retrying in {delay} seconds...")
#             time.sleep(delay)
#             delay *= 2
#         except Exception as e:
#             logging.error(f"Unexpected error: {e}. Retrying in {delay} seconds...")
#             time.sleep(delay)
#             delay *= 2
#     raise Exception("Max retries exceeded")

# def extract_text_from_docx(docx_stream: io.BytesIO) -> str:
#     logging.info("Extracting text from the DOCX.")
#     doc = Document(docx_stream)
#     text = ""
#     for para in doc.paragraphs:
#         text += para.text + '\n'
#     return text

# def extract_text_from_txt(txt_stream: io.BytesIO) -> str:
#     logging.info("Extracting text from the TXT.")
#     return txt_stream.read().decode('utf-8')

# def handle_file(file_stream: io.BytesIO, file_type: str) -> str:
#     if file_type == 'pdf':
#         return extract_text_from_pdf(file_stream)
#     elif file_type == 'docx':
#         return extract_text_from_docx(file_stream)
#     elif file_type == 'txt':
#         return extract_text_from_txt(file_stream)
#     else:
#         raise ValueError("Unsupported file format")

# @document_bp.route('/upload', methods=['POST'])
# def upload_file():
#     if 'file' in request.files:
#         file = request.files['file']
#         if file.filename == '':
#             return jsonify({'error': 'No selected file'}), 400
#         file_type = file.filename.rsplit('.', 1)[-1].lower()
#         if file_type not in ['pdf', 'docx', 'txt']:
#             return jsonify({'error': 'Unsupported file format'}), 400
#         file_stream = io.BytesIO(file.read())
#     elif 'url' in request.json:
#         url = request.json['url']
#         if not url:
#             return jsonify({'error': 'No URL provided'}), 400
#         try:
#             response = requests.get(url, stream=True)
#             response.raise_for_status()
#             content_type = response.headers.get('Content-Type', '')
#             if 'application/pdf' in content_type:
#                 file_type = 'pdf'
#             elif 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' in content_type:
#                 file_type = 'docx'
#             elif 'text/plain' in content_type:
#                 file_type = 'txt'
#             else:
#                 return jsonify({'error': 'Unsupported file format from URL'}), 400
#             file_stream = io.BytesIO(response.content)
#         except requests.RequestException as e:
#             logging.error(f"Error fetching file from URL: {e}")
#             return jsonify({'error': 'Error fetching file from URL'}), 500
#     else:
#         return jsonify({'error': 'No file or URL provided'}), 400

#     try:
#         text = handle_file(file_stream, file_type)
#         chunks = chunk_text(text)
        
#         llm = genai.GenerativeModel("gemini-pro")
#         summaries = []

#         logging.info("Starting the summarization process.")
#         for i, chunk in enumerate(chunks):
#             try:
#                 summary = summarize_chunk(llm, chunk)
#                 if summary:
#                     summaries.append(summary)
#                     logging.info(f"Chunk {i+1}/{len(chunks)} summarized successfully.")
#                 else:
#                     logging.warning(f"Chunk {i+1} could not be summarized. Skipping.")
#                 time.sleep(1)
#             except Exception as e:
#                 logging.error(f"Error processing chunk {i+1}: {e}")
#                 continue

#         final_summary = " ".join(filter(None, summaries))
#         return jsonify({'summary': final_summary})
#     except Exception as e:
#         logging.error(f"An error occurred: {e}")
#         return jsonify({'error': str(e)}), 500


# //////////////////////////////////////////////////
# response in jsonify
# //////////////////////////////////////////////////


from flask import Blueprint, request, jsonify
import io
import os
import logging
import requests
import PyPDF2
import unicodedata
from docx import Document
import google.generativeai as genai
import time
from google.api_core.exceptions import ResourceExhausted

# Blueprint for document routes
document_bp = Blueprint('document_bp', __name__)

# Load environment variables
api_key = os.getenv("GOOGLE_API_KEY")
genai.configure(api_key=api_key)

def extract_text_from_pdf(pdf_stream: io.BytesIO) -> str:
    logging.info("Extracting text from the PDF.")
    pdf_reader = PyPDF2.PdfReader(pdf_stream)
    text = ""
    for page in pdf_reader.pages:
        page_text = page.extract_text() or ""
        text += page_text
    return text

def normalize_text(text: str) -> str:
    return unicodedata.normalize('NFKC', text)

def chunk_text(text: str, chunk_size: int = 5000) -> list:
    text = normalize_text(text)
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        if end < len(text):
            while end < len(text) and text[end] in ('\uD800', '\uDBFF'):
                end += 1
        chunks.append(text[start:end])
        start = end
    return chunks

def summarize_chunk(llm, chunk: str, retries: int = 3, delay: int = 1) -> str:
    prompt = f"Please provide a well-defined summary of the following text:\n TEXT: {chunk}"
    for attempt in range(retries):
        try:
            response = llm.generate_content(prompt)
            summary_text = response.text
            return summary_text
        except ResourceExhausted as e:
            logging.warning(f"Resource exhausted error: {e}. Retrying in {delay} seconds...")
            time.sleep(delay)
            delay *= 2
        except Exception as e:
            logging.error(f"Unexpected error: {e}. Retrying in {delay} seconds...")
            time.sleep(delay)
            delay *= 2
    raise Exception("Max retries exceeded")

def extract_text_from_docx(docx_stream: io.BytesIO) -> str:
    logging.info("Extracting text from the DOCX.")
    doc = Document(docx_stream)
    text = ""
    for para in doc.paragraphs:
        text += para.text + '\n'
    return text

def extract_text_from_txt(txt_stream: io.BytesIO) -> str:
    logging.info("Extracting text from the TXT.")
    return txt_stream.read().decode('utf-8')

def handle_file(file_stream: io.BytesIO, file_type: str) -> str:
    if file_type == 'pdf':
        return extract_text_from_pdf(file_stream)
    elif file_type == 'docx':
        return extract_text_from_docx(file_stream)
    elif file_type == 'txt':
        return extract_text_from_txt(file_stream)
    else:
        raise ValueError("Unsupported file format")

@document_bp.route('/upload', methods=['POST'])
def upload_file():
    if 'file' in request.files:
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'No selected file'}), 400
        file_type = file.filename.rsplit('.', 1)[-1].lower()
        if file_type not in ['pdf', 'docx', 'txt']:
            return jsonify({'error': 'Unsupported file format'}), 400
        file_stream = io.BytesIO(file.read())
    elif 'url' in request.json:
        url = request.json['url']
        if not url:
            return jsonify({'error': 'No URL provided'}), 400
        try:
            response = requests.get(url, stream=True)
            response.raise_for_status()
            content_type = response.headers.get('Content-Type', '')
            if 'application/pdf' in content_type:
                file_type = 'pdf'
            elif 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' in content_type:
                file_type = 'docx'
            elif 'text/plain' in content_type:
                file_type = 'txt'
            else:
                return jsonify({'error': 'Unsupported file format from URL'}), 400
            file_stream = io.BytesIO(response.content)
        except requests.RequestException as e:
            logging.error(f"Error fetching file from URL: {e}")
            return jsonify({'error': 'Error fetching file from URL'}), 500
    else:
        return jsonify({'error': 'No file or URL provided'}), 400

    try:
        text = handle_file(file_stream, file_type)
        chunks = chunk_text(text)
        
        llm = genai.GenerativeModel("gemini-pro")
        summaries = []

        logging.info("Starting the summarization process.")
        for i, chunk in enumerate(chunks):
            try:
                summary = summarize_chunk(llm, chunk)
                if summary:
                    summaries.append(summary)
                    logging.info(f"Chunk {i+1}/{len(chunks)} summarized successfully.")
                else:
                    logging.warning(f"Chunk {i+1} could not be summarized. Skipping.")
                time.sleep(1)
            except Exception as e:
                logging.error(f"Error processing chunk {i+1}: {e}")
                continue

        final_summary = " ".join(filter(None, summaries))

        # Generating heading and detailed paragraphs
        heading_prompt = f"Provide a concise heading for the following text:\n\nTEXT: {final_summary}"
        paragraphs_prompt = f"Provide detailed paragraphs summarizing the following text:\n\nTEXT: {final_summary}"

        try:
            heading_response = llm.generate_content(heading_prompt)
            heading = heading_response.text.strip()
        except Exception as e:
            logging.error(f"Error generating heading: {e}")
            heading = "Summary"

        try:
            paragraphs_response = llm.generate_content(paragraphs_prompt)
            paragraphs = paragraphs_response.text.strip()
        except Exception as e:
            logging.error(f"Error generating paragraphs: {e}")
            paragraphs = final_summary

        return jsonify({'heading': heading, 'paragraphs': paragraphs})
    except Exception as e:
        logging.error(f"An error occurred: {e}")
        return jsonify({'error': str(e)}), 500
